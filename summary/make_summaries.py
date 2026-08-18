from docx import Document
from docx.shared import Pt, Inches, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x59, 0x59, 0x59)
OUT = "/sessions/blissful-zen-volta/mnt/earlyaijobs/summary/"


def new_doc():
    doc = Document()
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    for sec in doc.sections:
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
        sec.left_margin = sec.right_margin = Inches(0.9)
        sec.top_margin = sec.bottom_margin = Inches(0.8)
    return doc


def title(doc, text, sub):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(24), True, NAVY
    p2 = doc.add_paragraph(); r2 = p2.add_run(sub)
    r2.font.size, r2.font.italic, r2.font.color.rgb = Pt(11.5), True, GRAY
    p2.paragraph_format.space_after = Pt(14)


def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size, r.font.bold, r.font.color.rgb = Pt(14), True, NAVY
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)


def bullet(doc, text, lead=None):
    p = doc.add_paragraph(style="List Bullet")
    if lead:
        r = p.add_run(lead + " — "); r.font.bold = True
    p.add_run(text); p.paragraph_format.space_after = Pt(4)


# ============================================================
# DAY 3 — CLASSIFICATION
# ============================================================
d = new_doc()
title(d, "EarlyAIJobs — Build Journal",
      "Day 3: Classification — Teaching the System What Each Job Is  •  August 7–15, 2026")

h1(d, "What We Accomplished")
bullet(d, "Defined the product properly: EarlyAIJobs lists EVERY open job at approved AI companies. The classifier describes roles; it does not decide who gets listed. A recruiter at Anthropic belongs on the site.")
bullet(d, "Built and then deliberately retired an LLM classifier (six versions) that called Claude's API to categorise jobs, with validation, corrective retries, QA flags and confidence scores.")
bullet(d, "Replaced it with a deterministic classifier — classify-simple.js — using title rules first, description rules second, and 'other' as a legitimate answer. No AI calls, free, instant, auditable.")
bullet(d, "Locked an 18-category taxonomy: engineering, research, data, product, design, infrastructure, security, solutions, sales, marketing, customer-success, operations, legal-compliance, policy, people, finance, education, other.")
bullet(d, "Discovered and fixed a critical evidence bug: job descriptions were truncated at 800 characters, which is entirely company boilerplate. The classifier had almost no real evidence to work with.")
bullet(d, "Removed truncation entirely — full descriptions now stored — and preserved HTML structure so section headings survive.")
bullet(d, "Trimmed the company list from 33 to the 6 approved AI companies, and deleted 4,936 rows from de-scoped companies after verifying the count first.")
bullet(d, "Classified all 2,493 jobs. 'Other' rate: 3.4%.")

h1(d, "Concepts Learned (in plain English)")
bullet(d, "Written instructions given to an AI. Writing one is management, not programming: role, rules, edge cases, output format. It is an SOP for a very fast, very literal intern.", "Prompt")
bullet(d, "Code that checks the AI's answer before saving it. If the model returns a category that does not exist, the answer is rejected — not stored.", "Validation")
bullet(d, "When validation fails, show the model its own error and ask again. Far more effective than repeating the same question.", "Corrective retry")
bullet(d, "The model's self-reported certainty. NOT a measurement — it is a number generated the same way as any other word. Every misclassification we caught arrived at 90%+ confidence.", "Confidence score")
bullet(d, "Rules written in code rather than judgment made by a model. Same input always gives same output, every decision is auditable, and re-running is free.", "Deterministic classification")
bullet(d, "When two title rules both match, the longer phrase wins. 'machine learning engineer' beats 'engineer'; 'infrastructure financing' beats any technical infrastructure phrase.", "Longest-match wins")
bullet(d, "Repeated company text ('Anthropic is an AI safety and research company') that appears in every posting. It must never classify a job — it was dragging unrelated roles into 'research'.", "Company boilerplate")
bullet(d, "A rule requiring role-level evidence before assigning a category. Research needs research responsibilities, not a company description mentioning research.", "Eligibility gate")
bullet(d, "Only classify jobs that have never been categorised. Prevents re-labelling thousands of unchanged rows on every scheduled run.", "--only-new")

h1(d, "Judgment Lessons (the valuable part)")
bullet(d, "We spent hours tuning rules against input that was 90% boilerplate. Diagnose your data before you tune your logic.", "Verify inputs before tuning logic")
bullet(d, "A model that misreads a job misreads it fluently. A metric that cannot detect its own blind spots is decoration.", "Confidently wrong beats uncertain")
bullet(d, "Not every job maps cleanly to a category. Forcing 'Head of Music' somewhere is worse than admitting it does not fit.", "'Other' is a legitimate answer")
bullet(d, "Fix patterns, not individual jobs. Ten 'Safeguards Enforcement Analyst' roles in 'other' is a rule gap; one 'Power Trading Lead' is not.", "Fix systemic, ignore one-offs")
bullet(d, "Claude built, GPT reviewed, and the reviews caught real bugs. But by round six we were refining a classifier for a website that did not exist. Review is valuable; unbounded review is procrastination.", "Two AIs reviewing each other")

h1(d, "The Numbers")
bullet(d, "2,493 jobs classified  •  18 categories  •  3.4% 'other'  •  6 approved companies  •  $0 running cost  •  classifier versions built: 7")
bullet(d, "Distribution: engineering 27.6% · solutions 16.4% · sales 14.8% · operations 8.1% · marketing 4.7% · research 4.1% · other 3.4% · remaining 11 categories 20.9%")

h1(d, "Mistakes That Taught Us")
bullet(d, "Descriptions truncated at 800 characters, so the classifier read company boilerplate instead of responsibilities. Cost: most of two days of misdirected tuning.")
bullet(d, "Trusted model-reported confidence as a quality signal. Every wrong answer scored 90%+.")
bullet(d, "Ran SQL against the wrong Supabase project (there were two with similar names). The habit of running SELECT COUNT before DELETE caught it.")
bullet(d, "Pasted an API token at a shell prompt instead of a password prompt. Anything displayed anywhere is burned — rotate it.")

h1(d, "What's Next")
bullet(d, "Website build: Next.js front end reading from Supabase.")
bullet(d, "Deployment to DigitalOcean App Platform with a custom domain.")
bullet(d, "Automated hourly pipeline so the dataset maintains itself.")

d.save(OUT + "EarlyAIJobs-Day3-Summary.docx")


# ============================================================
# DAYS 4–6 — WEBSITE, DEPLOYMENT, AUTOMATION
# ============================================================
d = new_doc()
title(d, "EarlyAIJobs — Build Journal",
      "Days 4–6: Website, Deployment & Automation  •  August 16–17, 2026")

h1(d, "What We Accomplished")
bullet(d, "Built the website with Next.js 15: homepage with company/category/remote filters and live counts, title search, 50-per-page pagination, freshness badges under 48 hours, and individual job pages with outbound apply links.")
bullet(d, "Added SEO from day one: sitemap with ~2,520 URLs, robots.txt, per-job titles and descriptions, and JobPosting structured data with directApply set to false (we link out — misdeclaring that is penalised).")
bullet(d, "Chose DigitalOcean App Platform over Vercel for one concrete reason: Vercel's free tier caps scheduled tasks at once per day, which contradicts a product called EarlyAIJobs.")
bullet(d, "Pushed the repo to GitHub and deployed. Site live on a *.ondigitalocean.app URL within minutes.")
bullet(d, "Moved DNS from GoDaddy to DigitalOcean nameservers, attached earlyaijobs.com, HTTPS certificate issued automatically.")
bullet(d, "Applied a colour palette drawn from a painting: deep indigo for authority (header, buttons, active filters) and sage green for freshness (badges, hover states), with a near-white reading surface so job titles stay the highest-contrast element.")
bullet(d, "Hardened the collector against unattended operation, then created a scheduled Job component running hourly.")
bullet(d, "Added expired-job pages that stay live with a 'no longer accepting applications' notice and drop their JobPosting structured data.")

h1(d, "Concepts Learned (in plain English)")
bullet(d, "The framework our website is built with. Server components render pages on the server, so database credentials never reach the browser.", "Next.js")
bullet(d, "How long a rendered page is cached before being regenerated. Ours is 300 seconds.", "revalidate")
bullet(d, "Putting code on computers that serve it to the public. Push to GitHub → the platform builds → live in minutes.", "Deploying")
bullet(d, "The phone book of the internet. Nameservers decide who answers questions about your domain; records say where each name points.", "DNS / nameservers")
bullet(d, "The apex is earlyaijobs.com; www is a subdomain. www must be a RECORD inside the apex zone, never its own separate zone — that mistake cost us an afternoon.", "Apex vs subdomain")
bullet(d, "A list of every URL worth crawling, plus a file saying what crawlers may access.", "Sitemap / robots.txt")
bullet(d, "Hidden machine-readable facts about a page. For jobs it is what makes a listing eligible for Google's jobs experience.", "Structured data (JSON-LD)")
bullet(d, "The preview card shown when a link is shared. Disproportionately important on launch day.", "OG image")
bullet(d, "The platform's automatic build process. It runs BEFORE any custom build command, which is why 'just set the build command' often does not do what you expect.", "Buildpack")
bullet(d, "Runs on a schedule and exits, billed only for runtime. A worker runs continuously and is billed always. For a 40-second task every hour, the job is far cheaper.", "Scheduled job vs worker")
bullet(d, "Minute-hour-day-month-weekday. '0 * * * *' means minute 0 of every hour.", "Cron expression")
bullet(d, "0 means success, anything else means failure. It is how platforms detect that a scheduled run went wrong.", "Exit code")

h1(d, "Production Safety (the most important work of these days)")
bullet(d, "The original collector ended with one global sweep marking every job not seen this run as closed. If one platform had an outage, every job from that platform would be closed and the site would empty out.", "The hazard we found")
bullet(d, "Jobs are closed only for a company whose fetch succeeded in that run. A global sweep is now impossible.", "Fix 1 — per-company reconciliation")
bullet(d, "Fetch error, missing feed, or a failed database write means that company is skipped entirely and reported. Its existing jobs are untouched.", "Fix 2 — failed feeds skipped")
bullet(d, "A successful response containing zero jobs is ambiguous, not authoritative. Never treated as 'everything closed'.", "Fix 3 — empty feeds skipped")
bullet(d, "If a successful fetch would close more than 40% of a company's open jobs, the script refuses and logs loudly for human review.", "Fix 4 — 40% close ceiling")
bullet(d, "Closure is the LAST mutating step, reached only after a successful fetch AND a fully successful write.", "Fix 5 — operation order")
bullet(d, "Within an hour of shipping, OpenAI's writes failed on a network error. The old code would have closed 746 jobs. The new code reported WRITE-FAILED, left the data intact, and named the cause.", "It paid off immediately")

h1(d, "Mistakes That Taught Us")
bullet(d, "Sidebar counts capped at 1,000 while the header said 2,493. Supabase caps any single response at 1,000 rows, so counting by fetching rows silently under-reports. Fix: ask the database to count. Caught by noticing two numbers on one screen that should have agreed.")
bullet(d, "Upsert batches of 500 became multi-megabyte payloads once full descriptions were stored, causing write failures. Batch size is a function of payload size, not row count. Fixed with 100-row batches plus retries.")
bullet(d, "Created www.earlyaijobs.com as a separate DNS zone instead of a record inside the apex zone. A standalone subdomain zone is never delegated to, so it cannot resolve.")
bullet(d, "The scheduled job's first build failed on /sitemap.xml. Root cause was deeper: lib/db.js created the Supabase client at import time and threw without credentials. Fixed by making the build resilient rather than copying public keys onto the job — which also protects the website's build from a future outage.")
bullet(d, "A leftover .git/index.lock blocked commits. Created by a read-only Git command from a sandbox that could not clean up after itself.")

h1(d, "The Numbers")
bullet(d, "2,493 open jobs live  •  6 companies  •  18 categories  •  ~2,520 sitemap URLs  •  hourly refresh  •  pipeline runtime 0.6 min  •  ~$5–7/month total cost")

h1(d, "What's Next")
bullet(d, "Verify the scheduler: manual run, log inspection, before/after counts, then observe one automatic hourly run.")
bullet(d, "Expired-listing retention window (closed jobs are already hidden from all site queries; the window only affects direct URLs).")
bullet(d, "OG image and favicon before the LinkedIn launch post.")
bullet(d, "An /about page, company logos, and empty-state copy.")

d.save(OUT + "EarlyAIJobs-Days4-6-Summary.docx")
print("written both")
